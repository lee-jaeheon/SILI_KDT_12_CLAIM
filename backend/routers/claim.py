import uuid
import shutil
from io import BytesIO
from pathlib import Path
from typing import Optional
from urllib.parse import quote

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image as PILImage
from fastapi import APIRouter, Body, Depends, File, Form, HTTPException, Query, UploadFile
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from backend.core.uploads import check_size, MAX_IMAGE_MB
from backend.routers.auth import get_current_user
from backend.models.database import (
    insert_report, get_report, list_reports, update_report, delete_report,
    insert_image, get_images, delete_image,
    search_similar, get_defect_types, ReportNotFoundError,
)

router = APIRouter(prefix="/api/reports", tags=["reports"])

UPLOAD_DIR = Path(__file__).parent.parent.parent / "uploads"
UPLOAD_DIR.mkdir(exist_ok=True)

ALLOWED_EXTENSIONS = {".jpg", ".jpeg", ".png", ".webp", ".bmp"}

_TEMPLATE_PATH = Path(__file__).parent.parent / "부적합_처리_보고서_양식.docx"

_DEFECT_CHECKBOX = {
    "OUTER_DAMAGE": "외관손상",
    "SEALING":      "실링 불량",
    "HEMMING":      "헤밍 불량",
    "HOLE_DEFORM":  "홀 변형",
}


_TWIPS_PER_CM = 566.93


def _get_row_height_twips(row_obj) -> int:
    trPr = row_obj._tr.find(qn("w:trPr"))
    if trPr is None:
        return 0
    trH = trPr.find(qn("w:trHeight"))
    if trH is None:
        return 0
    return int(trH.get(qn("w:val"), "0"))


def _set_row_exact_height(row_obj, twips: int):
    tr = row_obj._tr
    trPr = tr.find(qn("w:trPr"))
    if trPr is None:
        trPr = OxmlElement("w:trPr")
        tr.insert(0, trPr)
    trH = trPr.find(qn("w:trHeight"))
    if trH is None:
        trH = OxmlElement("w:trHeight")
        trPr.append(trH)
    trH.set(qn("w:val"), str(twips))
    trH.set(qn("w:hRule"), "exact")


def _fill_cell(table, row: int, col: int, text: str, font_size: int = 10):
    """원본 단락 XML을 유지하면서 텍스트 run만 교체. spacing을 0으로 강제해 셀 팽창 방지."""
    cell = table.cell(row, col)
    tc = cell._tc
    p = tc.find(qn("w:p"))
    if p is None:
        p = OxmlElement("w:p")
        tc.append(p)
    for r in p.findall(qn("w:r")):
        p.remove(r)
    for extra_p in tc.findall(qn("w:p"))[1:]:
        tc.remove(extra_p)

    # pPr 안의 spacing을 before=0/after=0으로 강제 → 칸 팽창 방지
    pPr = p.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p.insert(0, pPr)
    sp = pPr.find(qn("w:spacing"))
    if sp is None:
        sp = OxmlElement("w:spacing")
        pPr.append(sp)
    sp.set(qn("w:before"), "0")
    sp.set(qn("w:after"), "0")

    r_el = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    for tag in ("w:sz", "w:szCs"):
        sz = OxmlElement(tag)
        sz.set(qn("w:val"), str(font_size * 2))
        rPr.append(sz)
    r_el.append(rPr)
    t_el = OxmlElement("w:t")
    t_el.text = str(text or "")
    if t_el.text:
        t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r_el.append(t_el)
    p.append(r_el)



def _generate_ncr_docx(report: dict) -> BytesIO:
    doc = Document(str(_TEMPLATE_PATH))
    tbl = doc.tables[0]

    # tblLayout fixed → Word가 내용 기준으로 표를 재계산하지 않도록 고정
    tblPr = tbl._tbl.find(qn("w:tblPr"))
    tblLayout = tblPr.find(qn("w:tblLayout"))
    if tblLayout is None:
        tblLayout = OxmlElement("w:tblLayout")
        tblPr.append(tblLayout)
    tblLayout.set(qn("w:type"), "fixed")

    dq  = report.get("delivery_quantity")
    dfq = report.get("defect_quantity")

    # 문서번호
    _fill_cell(tbl, 1, 9, report.get("document_no") or "")

    # 발행일 / 고객사
    _fill_cell(tbl, 2, 4, str(report.get("issue_date")    or "")[:10])
    _fill_cell(tbl, 2, 8, report.get("customer_name")     or "")

    # 납품일 / 납품수량
    _fill_cell(tbl, 3, 4, str(report.get("delivery_date") or "")[:10])
    _fill_cell(tbl, 3, 8, f"{dq} EA" if dq else "")

    # 품명/품번 / 불량수량
    product = " / ".join(filter(None, [report.get("product_name"), report.get("product_no")]))
    _fill_cell(tbl, 4, 4, product)
    _fill_cell(tbl, 4, 8, f"{dfq} EA" if dfq else "")

    # 불량유형 체크박스
    defect_type = report.get("defect_type") or ""
    cell = tbl.cell(5, 4)
    tc5 = cell._tc
    p5 = tc5.find(qn("w:p"))
    if p5 is None:
        p5 = OxmlElement("w:p"); tc5.append(p5)
    for r in p5.findall(qn("w:r")): p5.remove(r)
    for code, label in _DEFECT_CHECKBOX.items():
        mark = "☑" if code == defect_type else "☐"
        r_el = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        for tag in ("w:sz", "w:szCs"):
            sz = OxmlElement(tag); sz.set(qn("w:val"), "20"); rPr.append(sz)
        r_el.append(rPr)
        t_el = OxmlElement("w:t")
        t_el.text = f"{mark} {label}    "
        t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r_el.append(t_el); p5.append(r_el)

    # 불량 사진 — 원본 템플릿 이미지 칸(7.05cm) 안에 비율 유지하며 삽입
    images = report.get("images") or []
    if images:
        img_path = Path(__file__).parent.parent.parent / images[0]["image_path"].lstrip("/")
        if img_path.exists():
            cell = tbl.cell(7, 1)
            tc7 = cell._tc
            p7 = tc7.find(qn("w:p"))
            if p7 is None:
                p7 = OxmlElement("w:p"); tc7.append(p7)
            for r in p7.findall(qn("w:r")): p7.remove(r)
            for extra in tc7.findall(qn("w:p"))[1:]: tc7.remove(extra)
            # 이미지 셀 spacing도 0으로 강제
            pPr7 = p7.find(qn("w:pPr"))
            if pPr7 is None:
                pPr7 = OxmlElement("w:pPr"); p7.insert(0, pPr7)
            sp7 = pPr7.find(qn("w:spacing"))
            if sp7 is None:
                sp7 = OxmlElement("w:spacing"); pPr7.append(sp7)
            sp7.set(qn("w:before"), "0")
            sp7.set(qn("w:after"), "0")
            run = cell.paragraphs[0].add_run()
            img_buf, img_w, img_h = _fit_image_in_cell(img_path, max_w_cm=14.5, max_h_cm=6.7)
            run.add_picture(img_buf, width=img_w, height=img_h)

    # 원인분석 / 시정조치 / 예방대책
    _fill_cell(tbl, 8,  2, report.get("root_cause_analysis") or "")
    _fill_cell(tbl, 9,  2, report.get("corrective_action")   or "")
    _fill_cell(tbl, 10, 2, report.get("preventive_action")   or "")

    # 서명란 (작성 / 검토 / 승인)
    _fill_cell(tbl, 12, 0, report.get("author_name")   or "")
    _fill_cell(tbl, 12, 5, report.get("reviewer_name") or "")
    _fill_cell(tbl, 12, 7, report.get("approver_name") or "")

    # ── 전체 행 높이를 exact로 고정 ──
    # 13697(24.16cm)보다 약간 작게 잡아 Word 테두리·여백·줄높이 오차를 흡수하면서
    # 페이지 하단이 과하게 남지 않도록 한다.
    _PAGE_TWIPS = 13600  # ≈ 23.99cm (여유 ~0.17cm)
    _ROW_H = {
        0:  650,   # 로고/문서번호 상단
        1:  900,   # 로고가 잘리지 않도록 상단 병합 셀 높이 확보
        2:  350,   # 발행일 / 고객사 (1줄)
        3:  350,   # 납품일 / 납품수량 (1줄)
        4:  650,   # 품명/품번 (2줄 허용)
        5:  350,   # 불량유형 체크박스 (1줄)
        6:  350,   # 전체사진 라벨 (1줄)
        # 7 : 아래에서 나머지로 계산
        8:  1694,  # 원인분석 (원본 그대로)
        9:  1694,  # 시정조치 (원본 그대로)
        10: 1694,  # 예방대책 (원본 그대로)
        11:  413,  # 서명 라벨 (원본 그대로)
        12:  547,  # 이름란 (원본 그대로)
    }
    _ROW_H[7] = _PAGE_TWIPS - sum(_ROW_H.values())  # 이미지 행이 나머지 공간 전부 흡수
    for idx, h in _ROW_H.items():
        _set_row_exact_height(tbl.rows[idx], h)

    # 테이블 뒤 빈 단락의 기본 after-spacing(160 twips) 제거 → 2페이지 밀림 방지
    for p in doc.element.body.findall(qn("w:p")):
        pPr = p.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            p.insert(0, pPr)
        sp = pPr.find(qn("w:spacing"))
        if sp is None:
            sp = OxmlElement("w:spacing")
            pPr.append(sp)
        sp.set(qn("w:before"), "0")
        sp.set(qn("w:after"), "0")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _fit_image_in_cell(img_path: Path, max_w_cm: float = 14.5, max_h_cm: float = 7.0):
    """이미지를 셀 박스에 비율 유지하며 맞춤. (buf, width_Cm, height_Cm) 반환"""
    img = PILImage.open(img_path)
    if img.mode not in ("RGB", "L"):
        img = img.convert("RGB")
    w_px, h_px = img.size
    aspect = w_px / h_px
    if aspect > max_w_cm / max_h_cm:  # 가로가 더 넓은 경우
        fit_w = max_w_cm
        fit_h = max_w_cm / aspect
    else:                              # 세로가 더 긴 경우
        fit_h = max_h_cm
        fit_w = max_h_cm * aspect
    # 150 DPI 기준으로 리사이즈 (파일 크기 최적화)
    t_w = int(fit_w / 2.54 * 150)
    t_h = int(fit_h / 2.54 * 150)
    img = img.resize((t_w, t_h), PILImage.LANCZOS)
    buf = BytesIO()
    img.save(buf, format="JPEG", quality=85)
    buf.seek(0)
    return buf, Cm(fit_w), Cm(fit_h)


def _save_image(image: UploadFile) -> str:
    ext = Path(image.filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail="지원하지 않는 이미지 형식입니다.")
    check_size(image, MAX_IMAGE_MB)
    filename  = f"{uuid.uuid4().hex}{ext}"
    save_path = UPLOAD_DIR / filename
    with open(save_path, "wb") as f:
        shutil.copyfileobj(image.file, f)
    return f"/uploads/{filename}"


def _delete_saved_image(image_path: Optional[str]):
    if not image_path:
        return
    full_path = Path(__file__).parent.parent.parent / image_path.lstrip("/")
    if full_path.exists():
        full_path.unlink()


# ── CRUD ───────────────────────────────────────────────────────────────────────

@router.get("/")
def get_reports(
    status: Optional[str] = Query(None),
    page:   int           = Query(1, ge=1),
    limit:  int           = Query(20, ge=1, le=500),
    _: dict = Depends(get_current_user),
):
    offset = (page - 1) * limit
    result = list_reports(status=status, limit=limit, offset=offset)
    return {
        "total": result["total"],
        "page":  page,
        "limit": limit,
        "items": result["items"],
    }


@router.post("/")
async def create_report(
    customer_name:     str            = Form(...),
    defect_type:       str            = Form(""),
    defect_location:   str            = Form(""),
    product_name:      str            = Form(""),
    product_no:        str            = Form(""),
    part_name:         str            = Form(""),
    process_name:      str            = Form(""),
    lot_no:            str            = Form(""),
    delivery_quantity: Optional[int]  = Form(None),
    defect_quantity:   Optional[int]  = Form(None),
    handler:           str            = Form(""),
    author_name:       str            = Form(""),
    delivery_date:     str            = Form(""),
    issue_date:        str            = Form(""),
    defect_code:         str            = Form(""),
    ai_defect_type:      str            = Form(""),
    ai_confidence:       Optional[float] = Form(None),
    claim_text:          str            = Form(""),
    extracted_text:      str            = Form(""),
    root_cause_analysis: str            = Form(""),
    corrective_action:   str            = Form(""),
    preventive_action:   str            = Form(""),
    reviewer_name:       str            = Form(""),
    approver_name:       str            = Form(""),
    image: Optional[UploadFile] = File(None),
    _: dict = Depends(get_current_user),
):
    # defect_type 유효성 검증
    if defect_type:
        valid_codes = {d["code"] for d in get_defect_types()}
        if defect_type not in valid_codes:
            raise HTTPException(status_code=400, detail=f"유효하지 않은 불량 유형입니다: {defect_type}")

    image_path = None
    if image and image.filename:
        image_path = _save_image(image)

    report_id = None
    try:
        report_id = insert_report(
            customer_name=customer_name,
            defect_type=defect_type or None,
            defect_code=defect_code or None,
            defect_location=defect_location or None,
            ai_defect_type=ai_defect_type or None,
            ai_confidence=ai_confidence,
            product_name=product_name or None,
            product_no=product_no or None,
            part_name=part_name or None,
            process_name=process_name or None,
            lot_no=lot_no or None,
            delivery_quantity=delivery_quantity,
            defect_quantity=defect_quantity,
            handler=handler or None,
            author_name=author_name or None,
            delivery_date=delivery_date or None,
            issue_date=issue_date or None,
            claim_text=claim_text or None,
            extracted_text=extracted_text or None,
            root_cause_analysis=root_cause_analysis or None,
            corrective_action=corrective_action or None,
            preventive_action=preventive_action or None,
            reviewer_name=reviewer_name or None,
            approver_name=approver_name or None,
        )

        if image_path:
            insert_image(report_id, image_path, image_type="불량부위")
    except Exception:
        _delete_saved_image(image_path)
        if report_id is not None:
            try:
                delete_report(report_id)
            except Exception:
                pass
        raise

    return {"report_id": report_id, "message": "보고서가 생성되었습니다.", "image_path": image_path}


@router.get("/similar")
def get_similar(
    defect_type: Optional[str] = Query(None),
    customer_name: Optional[str] = Query(None),
    product_name: Optional[str] = Query(None),
    product_no: Optional[str] = Query(None),
    part_name: Optional[str] = Query(None),
    process_name: Optional[str] = Query(None),
    defect_location: Optional[str] = Query(None),
    claim_text: Optional[str] = Query(None),
    extracted_text: Optional[str] = Query(None),
    claim_summary: Optional[str] = Query(None),
    root_cause_analysis: Optional[str] = Query(None),
    corrective_action: Optional[str] = Query(None),
    preventive_action: Optional[str] = Query(None),
    limit: int = Query(5, ge=1, le=20),
    min_score: int = Query(60, ge=0, le=100),
    _: dict = Depends(get_current_user),
):
    return search_similar(
        defect_type=defect_type,
        customer_name=customer_name,
        product_name=product_name,
        product_no=product_no,
        part_name=part_name,
        process_name=process_name,
        defect_location=defect_location,
        claim_text=claim_text,
        extracted_text=extracted_text,
        claim_summary=claim_summary,
        root_cause_analysis=root_cause_analysis,
        corrective_action=corrective_action,
        preventive_action=preventive_action,
        limit=limit,
        min_score=min_score,
    )


@router.get("/{report_id}/preview")
def preview_report(report_id: int, _: dict = Depends(get_current_user)):
    report = get_report(report_id)
    if not report:
        raise HTTPException(status_code=404, detail="보고서를 찾을 수 없습니다.")
    buf = _generate_ncr_docx(report)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "inline"},
    )


@router.get("/{report_id}/download")
def download_report(report_id: int, _: dict = Depends(get_current_user)):
    report = get_report(report_id)
    if not report:
        raise HTTPException(status_code=404, detail="보고서를 찾을 수 없습니다.")
    buf      = _generate_ncr_docx(report)
    doc_no   = (report.get("document_no") or f"report_{report_id}").replace("/", "-")
    filename = quote(f"NCR_{doc_no}.docx", safe="")
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{filename}"},
    )


@router.get("/{report_id}")
def get_report_detail(report_id: int, _: dict = Depends(get_current_user)):
    row = get_report(report_id)
    if not row:
        raise HTTPException(status_code=404, detail="보고서를 찾을 수 없습니다.")
    return row


class UpdateBody(BaseModel):
    defect_type:         Optional[str]   = None
    defect_code:         Optional[str]   = None
    defect_location:     Optional[str]   = None
    product_name:        Optional[str]   = None
    product_no:          Optional[str]   = None
    part_name:           Optional[str]   = None
    process_name:        Optional[str]   = None
    lot_no:              Optional[str]   = None
    customer_name:       Optional[str]   = None
    delivery_quantity:   Optional[int]   = None
    defect_quantity:     Optional[int]   = None
    delivery_date:       Optional[str]   = None
    issue_date:          Optional[str]   = None
    root_cause_analysis: Optional[str]   = None
    corrective_action:   Optional[str]   = None
    preventive_action:   Optional[str]   = None
    handler:             Optional[str]   = None
    author_name:         Optional[str]   = None
    reviewer_name:       Optional[str]   = None
    approver_name:       Optional[str]   = None
    report_status:       Optional[str]   = None


@router.put("/{report_id}")
def update_report_api(report_id: int, body: UpdateBody, _: dict = Depends(get_current_user)):
    # defect_type이 들어오면 유효성 검증 (생성 API와 동일)
    if body.defect_type:
        valid_codes = {d["code"] for d in get_defect_types()}
        if body.defect_type not in valid_codes:
            raise HTTPException(status_code=400, detail=f"유효하지 않은 불량 유형입니다: {body.defect_type}")

    result = update_report(report_id, body.model_dump(exclude_none=True))
    if result == "missing":
        raise HTTPException(status_code=404, detail="보고서를 찾을 수 없습니다.")
    if result == "no_fields":
        raise HTTPException(status_code=400, detail="수정할 필드가 없습니다.")
    if result == "invalid_status":
        raise HTTPException(status_code=400, detail="유효하지 않은 상태값입니다.")
    return {"message": "수정되었습니다."}


@router.patch("/{report_id}/status")
def change_status(report_id: int, status: str = Body(..., embed=True), _: dict = Depends(get_current_user)):
    VALID_STATUSES = {"draft", "submitted", "approved", "rejected"}
    if status not in VALID_STATUSES:
        raise HTTPException(status_code=400, detail=f"유효하지 않은 상태입니다. 허용값: {', '.join(sorted(VALID_STATUSES))}")
    result = update_report(report_id, {"report_status": status})
    if result == "missing":
        raise HTTPException(status_code=404, detail="보고서를 찾을 수 없습니다.")
    return {"report_id": report_id, "report_status": status, "message": f"상태가 '{status}'로 변경되었습니다."}


@router.delete("/{report_id}")
def remove_report(report_id: int, _: dict = Depends(get_current_user)):
    if delete_report(report_id) == 0:
        raise HTTPException(status_code=404, detail="보고서를 찾을 수 없습니다.")
    return {"message": "삭제되었습니다."}


# ── 이미지 관리 ────────────────────────────────────────────────────────────────

@router.post("/{report_id}/images")
async def add_image(
    report_id: int,
    image: UploadFile = File(...),
    image_type: str = Form(""),
    image_description: str = Form(""),
    _: dict = Depends(get_current_user),
):
    image_path = _save_image(image)
    try:
        iid = insert_image(report_id, image_path, image_type or None, image_description or None)
    except ReportNotFoundError:
        _delete_saved_image(image_path)
        raise HTTPException(status_code=404, detail="보고서를 찾을 수 없습니다.")
    except Exception:
        _delete_saved_image(image_path)
        raise
    return {"image_id": iid, "image_path": image_path}


@router.get("/{report_id}/images")
def list_images(report_id: int, _: dict = Depends(get_current_user)):
    if get_report(report_id) is None:
        raise HTTPException(status_code=404, detail="보고서를 찾을 수 없습니다.")
    return get_images(report_id)


@router.delete("/images/{image_id}")
def remove_image(image_id: int, _: dict = Depends(get_current_user)):
    if delete_image(image_id) == 0:
        raise HTTPException(status_code=404, detail="이미지를 찾을 수 없습니다.")
    return {"message": "삭제되었습니다."}
