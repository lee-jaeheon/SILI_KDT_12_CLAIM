import asyncio
import io
import re
import json
import logging
import random

from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile
from fastapi.responses import StreamingResponse
from pathlib import Path

from backend.ai.ollama import call_ollama, stream_ollama
from backend.core.defects import DEFECT_CODES, label_of, part_label_of
from backend.core.uploads import read_with_limit, MAX_IMAGE_MB, MAX_DOC_MB
from backend.models.database import get_defect_types, search_similar
from backend.routers.auth import get_current_user

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api", tags=["ai"])

_model = None


def preload_model():
    """lifespan에서 서버 시작 시 YOLO 모델을 미리 로드."""
    global _model
    try:
        from backend.ai.classifier import load_model
        _model = load_model()
        logger.info("YOLO 모델 로드 완료")
    except Exception as e:
        logger.error(f"YOLO 모델 로드 실패: {e}")
        _model = False


def _get_model():
    global _model
    if _model is None:
        preload_model()
    return _model if _model else None


# ── 문서 내 이미지 추출 헬퍼 ──────────────────────────────────────────────────────

def _extract_first_image_pdf(data: bytes):
    """PDF에서 첫 번째 내장 이미지를 PNG bytes로 반환. 없거나 실패 시 None."""
    try:
        import io as _io
        from pdfminer.high_level import extract_pages
        from pdfminer.layout import LTImage
        from PIL import Image

        def _iter_images(layout):
            for el in layout:
                if isinstance(el, LTImage):
                    yield el
                elif hasattr(el, '__iter__'):
                    yield from _iter_images(el)

        for page_layout in extract_pages(_io.BytesIO(data)):
            for lt_img in _iter_images(page_layout):
                try:
                    raw = lt_img.stream.get_rawdata()
                    if not raw:
                        continue
                    img = Image.open(_io.BytesIO(raw)).convert("RGB")
                    buf = _io.BytesIO()
                    img.save(buf, format="PNG")
                    return buf.getvalue()
                except Exception:
                    continue
    except Exception as e:
        logger.warning("PDF 이미지 추출 실패: %s", e)
    return None


def _extract_first_image_docx(doc):
    """DOCX에서 첫 번째 내장 이미지를 PNG bytes로 반환. 없거나 실패 시 None."""
    try:
        import io as _io
        from PIL import Image
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    blob = rel.target_part.blob
                    img  = Image.open(_io.BytesIO(blob)).convert("RGB")
                    buf  = _io.BytesIO()
                    img.save(buf, format="PNG")
                    return buf.getvalue()
                except Exception:
                    continue
    except Exception as e:
        logger.warning("DOCX 이미지 추출 실패: %s", e)
    return None


def _classify_image_bytes(image_data: bytes):
    """이미지 bytes를 YOLO로 분류해 결과 dict 반환. 실패 시 None."""
    try:
        model    = _get_model()
        is_dummy = model is None

        if model:
            try:
                from backend.ai.classifier import predict
                defect_type, confidence = predict(model, image_data)
                if defect_type is None:
                    return {
                        "defect_type": None, "label": "미검출",
                        "confidence": 0.0, "confidence_pct": "0.0%",
                        "is_dummy": False, "is_detected": False,
                    }
            except Exception as e:
                logger.warning("문서 이미지 분류 실패 (더미 전환): %s", e)
                defect_type = random.choice(DEFECT_CODES)
                confidence  = round(random.uniform(0.55, 0.92), 3)
                is_dummy    = True
        else:
            defect_type = random.choice(DEFECT_CODES)
            confidence  = round(random.uniform(0.55, 0.92), 3)

        return {
            "defect_type":    defect_type,
            "label":          label_of(defect_type),
            "confidence":     confidence,
            "confidence_pct": f"{confidence * 100:.1f}%",
            "is_dummy":       is_dummy,
            "is_detected":    True,
        }
    except Exception as e:
        logger.warning("문서 이미지 AI 분류 실패: %s", e)
        return None


# ── AI 이미지 분류 ─────────────────────────────────────────────────────────────

def _run_classify(model, image_data: bytes) -> tuple:
    """YOLO 추론 동기 헬퍼. asyncio.to_thread용.
    Returns: (defect_type, confidence, is_detected, is_dummy)
    """
    if model:
        try:
            from backend.ai.classifier import predict
            defect_type, confidence = predict(model, image_data)
            if defect_type is None:
                return None, 0.0, False, False
        except Exception as e:
            logger.warning("모델 예측 실패 (더미 모드 전환): %s", e)
            return random.choice(DEFECT_CODES), round(random.uniform(0.55, 0.92), 3), True, True
        return defect_type, confidence, True, False
    else:
        return random.choice(DEFECT_CODES), round(random.uniform(0.55, 0.92), 3), True, True


@router.post("/classify")
async def classify_image(image: UploadFile = File(...), _: dict = Depends(get_current_user)):
    image_data = await read_with_limit(image, MAX_IMAGE_MB)
    model = _get_model()

    defect_type, confidence, is_detected, is_dummy = await asyncio.to_thread(
        _run_classify, model, image_data
    )

    if not is_detected:
        return {
            "defect_type":    None,
            "label":          "미검출",
            "confidence":     0.0,
            "confidence_pct": "0.0%",
            "is_dummy":       False,
            "is_detected":    False,
            "similar_cases":  [],
        }

    try:
        similar = await asyncio.to_thread(search_similar, defect_type=defect_type, limit=3)
    except Exception as e:
        logger.warning("유사사례 검색 실패 (분류 결과는 정상): %s", e)
        similar = []

    return {
        "defect_type":    defect_type,
        "label":          label_of(defect_type),
        "confidence":     confidence,
        "confidence_pct": f"{confidence * 100:.1f}%",
        "is_dummy":       is_dummy,
        "is_detected":    True,
        "similar_cases":  similar,
    }


# ── AI 보고서 스트리밍 생성 (SSE) ───────────────────────────────────────────────

@router.post("/generate-report-stream")
async def generate_report_stream(
    defect_type:     str   = Form(...),
    confidence:      float = Form(0.0),
    extracted_text:  str   = Form(""),
    customer_name:   str   = Form(""),
    product_name:    str   = Form(""),
    defect_location: str   = Form(""),
    part_category:   str   = Form("FRAME"),
    _: dict = Depends(get_current_user),
):
    defect_label = label_of(defect_type)
    part_label   = part_label_of(part_category)

    try:
        similar_cases = await asyncio.to_thread(
            search_similar,
            defect_type=defect_type,
            customer_name=customer_name or None,
            product_name=product_name or None,
            defect_location=defect_location or None,
            limit=3,
            min_score=30,
        )
    except Exception as e:
        logger.warning(f"유사사례 검색 실패 (보고서 생성은 계속): {e}")
        similar_cases = []

    similar_text = ""
    for i, case in enumerate(similar_cases, 1):
        similar_text += f"""
[유사 사례 {i}] 고객사: {case.get('customer_name','')} / 불량유형: {label_of(case.get('defect_type',''))}
  - 원인 분석: {case.get('root_cause_analysis') or '없음'}
  - 시정 조치: {case.get('corrective_action') or '없음'}
  - 재발 방지: {case.get('preventive_action') or '없음'}
""".strip() + "\n"

    if not similar_text:
        similar_text = "유사 사례 없음"

    prompt = f"""당신은 자동차 부품 품질보증 전문가입니다.
아래 정보를 바탕으로 부적합 보고서의 세 항목을 작성하라.
반드시 JSON만 출력하고 설명은 쓰지 마라.

[불량 정보]
- 부품 유형: {part_label}
- 불량 유형: {defect_label} (AI 신뢰도: {confidence*100:.1f}%)
- 고객사: {customer_name or '미입력'}
- 제품명: {product_name or '미입력'}
- 불량 위치: {defect_location or '미입력'}
- 문서 내용: {extracted_text[:2000] if extracted_text else '미입력'}

[유사 사례]
{similar_text}

출력 형식 (JSON):
{{"root_cause_analysis":"원인 분석 내용","corrective_action":"시정 조치 내용","preventive_action":"재발 방지 내용"}}

작성 기준:
- 유사 사례를 참고하되 현재 불량 정보에 맞게 구체적으로 작성할 것
- 각 항목은 2~4문장으로 작성할 것
- 한국어로 작성할 것
출력:"""

    def _dumps(obj):
        return json.dumps(obj, ensure_ascii=False, default=str)

    async def event_generator():
        yield f"data: {_dumps({'type': 'similar_cases', 'data': similar_cases})}\n\n"

        accumulated = ""
        try:
            async for token in stream_ollama(prompt):
                accumulated += token
                yield f"data: {_dumps({'type': 'token', 'text': token})}\n\n"
        except Exception as e:
            yield f"data: {_dumps({'type': 'error', 'detail': str(e)})}\n\n"
            return

        try:
            result = None
            try:
                result = json.loads(accumulated.strip())
            except json.JSONDecodeError:
                pass
            if result is None:
                match = re.search(r'\{.*\}', accumulated, re.DOTALL)
                if match:
                    try:
                        result = json.loads(match.group())
                    except json.JSONDecodeError:
                        pass
            if result is None:
                yield f"data: {_dumps({'type': 'error', 'detail': 'LLM 응답 파싱 실패'})}\n\n"
                return
            result["similar_cases"] = similar_cases
            yield f"data: {_dumps({'type': 'done', 'result': result})}\n\n"
        except Exception:
            yield f"data: {_dumps({'type': 'error', 'detail': 'JSON 파싱 실패'})}\n\n"

    return StreamingResponse(event_generator(), media_type="text/event-stream")


# ── 불량 유형 목록 ─────────────────────────────────────────────────────────────

@router.get("/defect-types")
def defect_types(_: dict = Depends(get_current_user)):
    return get_defect_types()


# ── 파일 텍스트 추출 ───────────────────────────────────────────────────────────

def _parse_file_sync(data: bytes, ext: str) -> tuple:
    """파일 bytes에서 텍스트와 첫 번째 이미지를 동기로 추출. asyncio.to_thread용."""
    text = ""
    image_data = None

    if ext == ".pdf":
        import pdfplumber, io as _io
        with pdfplumber.open(_io.BytesIO(data)) as pdf:
            text = "\n".join(p.extract_text() or "" for p in pdf.pages)
        image_data = _extract_first_image_pdf(data)
    elif ext == ".docx":
        from docx import Document
        import io as _io
        doc   = Document(_io.BytesIO(data))
        lines = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                lines.append(" | ".join(c.text.strip() for c in row.cells if c.text.strip()))
        text = "\n".join(lines)
        image_data = _extract_first_image_docx(doc)
    elif ext == ".eml":
        import email
        msg = email.message_from_bytes(data)
        for part in msg.walk():
            if part.get_content_type() == "text/plain":
                text += part.get_payload(decode=True).decode("utf-8", errors="ignore")
            elif part.get_content_type().startswith("image/") and image_data is None:
                image_data = part.get_payload(decode=True)
    else:
        raise ValueError("PDF, DOCX, EML 파일만 지원합니다.")

    return text, image_data


@router.post("/parse-file")
async def parse_file(file: UploadFile = File(...), _: dict = Depends(get_current_user)):
    data = await read_with_limit(file, MAX_DOC_MB)
    ext  = Path(file.filename).suffix.lower()

    if ext not in {".pdf", ".docx", ".eml"}:
        raise HTTPException(status_code=400, detail="PDF, DOCX, EML 파일만 지원합니다.")

    try:
        text, image_data = await asyncio.to_thread(_parse_file_sync, data, ext)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"파일 읽기 실패: {e}")

    if not text.strip():
        raise HTTPException(status_code=400, detail="파일에서 텍스트를 추출할 수 없습니다.")

    ai_result = await asyncio.to_thread(_classify_image_bytes, image_data) if image_data else None

    return {"extracted_text": text[:8000], "ai_result": ai_result}


# ── 텍스트 → 필드 파싱 (Ollama) ───────────────────────────────────────────────

@router.post("/parse-claim")
async def parse_claim(text: str = Form(...), _: dict = Depends(get_current_user)):
    prompt = f"""아래 클레임 텍스트에서 정보를 추출해서 JSON으로 출력하라. 반드시 JSON만 출력하고 설명은 쓰지 마라.

출력 형식:
{{"customer_name":"고객사명","product_name":"품명","product_no":"품번","defect_type":"OUTER_DAMAGE또는SEALING또는HEMMING또는HOLE_DEFORM또는GAP_DEFECT또는FASTENING_DEFECT","delivery_date":"YYYY-MM-DD형식날짜","delivery_quantity":수량숫자,"defect_quantity":불량수량숫자,"defect_location":"불량위치","handler":"담당자이름","root_cause_analysis":"원인분석내용","corrective_action":"시정조치내용","preventive_action":"재발방지내용"}}

정보가 없으면 null로 채워라. delivery_date는 반드시 YYYY-MM-DD 형식으로 변환해서 출력하라.

예시:
입력: "현대자동차 프레임 FRM-001 외관 손상 불량. 납품일 2024-03-15. 500개 중 47개 불량. 프레임 상단 접합부. 담당 이영수. 설비 충격으로 추정."
출력: {{"customer_name":"현대자동차","product_name":"프레임","product_no":"FRM-001","defect_type":"OUTER_DAMAGE","delivery_date":"2024-03-15","delivery_quantity":500,"defect_quantity":47,"defect_location":"프레임 상단 접합부","handler":"이영수","root_cause_analysis":"설비 충격으로 인한 외관 손상 추정","corrective_action":"설비 점검 및 충격 방지 조치","preventive_action":"정기 설비 점검 일정 수립"}}

클레임 텍스트:
{text}
출력:"""

    response = None
    try:
        response = await call_ollama(prompt)
        match = re.search(r'\{.*\}', response, re.DOTALL)
        if not match:
            logger.warning("parse_claim: LLM 응답에서 JSON 추출 실패. 응답: %s", response[:300])
            raise HTTPException(status_code=500, detail="파싱 실패")
        return json.loads(match.group())
    except HTTPException:
        raise
    except Exception as e:
        logger.warning("parse_claim: JSON 파싱 실패 (%s). 응답: %s", e, response[:300] if response else 'N/A')
        raise HTTPException(status_code=500, detail="JSON 파싱 실패")
